import requests 
from docx import Document
import os
from bs4 import BeautifulSoup
from config import SECRET_KEY_TRANSLATOR, SECRET_KEY_OPENAI,ENDPOINT_TRANSLATOR, ENDPOINT_OPENAI, TARGET_URL
from langchain_openai.chat_models.azure import AzureChatOpenAI

#############TRANSLATE API
subscription_key = SECRET_KEY_TRANSLATOR
endpoint = ENDPOINT_TRANSLATOR
location = "eastus"
target_language = "pt-br"

def translator_text(text: str, target_language: str):
    path = '/translate'
    constructed_url = endpoint + path

    headers = {
        'Ocp-Apim-Subscription-Key': subscription_key,
        # location required if you're using a multi-service or regional (not global) resource.
        'Ocp-Apim-Subscription-Region': location,
        'Content-type': 'application/json',
        'X-ClientTraceId': str(os.urandom(16))
    }

    # You can pass more than one object in body.
    body = [{
        'text': text
    }]

    params = {
                'api-version': '3.0',
                'from': 'en',
                'to': target_language
            }
    
    request = requests.post(constructed_url, params=params, headers=headers, json=body)
    response = request.json()
    return response[0]["translations"][0]["text"]

def translator_document(path: str):
    document = Document(path)
    full_text = []
    for paragraph in document.paragraphs:
        paragraph.text = translator_text(paragraph.text, target_language)
        full_text.append(paragraph.text)
    
    translated_doc = Document()
    for paragraph in full_text:
        translated_doc.add_paragraph(paragraph)

    path_translated_doc = path.replace(".docx", f"_translated_{target_language}.docx")
    translated_doc.save(path_translated_doc)

    return path_translated_doc


########################################


#############OPENAI API
client = AzureChatOpenAI(
    azure_endpoint=ENDPOINT_OPENAI, 
    api_key=SECRET_KEY_OPENAI, 
    api_version="2024-05-01-preview", 
    azure_deployment="gpt-4o-mini",
    max_retries=0, 
    max_tokens=200, 
    openai_api_type="azure"
)

def translate_article(text, lang):
    messages = [
        ("system", "Você atua como tradutor de textos"), 
        ("user", f"Traduza o {text} para o idioma {lang} e responda em markdown")
    ]

    response = client.invoke(messages)
    print(response.content)
    return response.content



def extract_text(url: str):
    response = requests.get(url)

    if response.status_code != 200:
        print(f"Failed to fetch the URL. Status code: {response.status_code}")
        return None

    soup = BeautifulSoup(response.text, "html.parser")
    text = soup.get_text().strip()
    return text

########################################

if __name__ == '__main__':
    try:
        # print(translator_document("musica.docx"))
        texto_url = extract_text(TARGET_URL)
        print(translate_article(texto_url, target_language))
    except Exception as e:
        print(e)

